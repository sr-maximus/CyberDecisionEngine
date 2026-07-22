from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "manual" / "Guia_Lectura_Informes_CyberDecisionEngine.docx"

NAVY = "17324D"
TEAL = "007F86"
CYAN = "13A8B0"
INK = "172433"
MUTED = "607286"
LIGHT = "EDF5F7"
PALE = "F5F8FA"
GOLD = "A66A00"
RED = "A33A3A"
GREEN = "19785B"
WHITE = "FFFFFF"
LINE = "CFDCE3"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run(run, size=11, color=INK, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def keep_table_rows_together(table):
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)


def set_paragraph_shading(paragraph, fill=LIGHT, border=TEAL):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    borders.append(left)
    p_pr.append(borders)


def add_field(paragraph, field):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, align=None, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_rich_para(doc, parts, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    for text, bold, color in parts:
        set_run(p.add_run(text), color=color, bold=bold)
    return p


def add_callout(doc, label, body, tone="info"):
    color = {"info": TEAL, "warning": GOLD, "danger": RED, "success": GREEN}[tone]
    fill = {"info": LIGHT, "warning": "FFF7E8", "danger": "FCEEEE", "success": "ECF7F2"}[tone]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(p, fill=fill, border=color)
    set_run(p.add_run(f"{label}: "), bold=True, color=color)
    set_run(p.add_run(body), color=INK)
    return p


def add_formula(doc, formula, explanation):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    set_paragraph_shading(p, fill="F3F6F8", border=NAVY)
    set_run(p.add_run(formula), size=12, color=NAVY, bold=False, font="Cambria Math")
    p2 = add_para(doc, explanation, size=9.5, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    return p, p2


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    set_run(p.add_run(text))
    return p


def start_steps(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    doc._cde_step_num_id = num_id


def add_step(doc, title, body):
    p = doc.add_paragraph()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(doc._cde_step_num_id))
    num_pr.extend([ilvl, num_id])
    p._p.get_or_add_pPr().append(num_pr)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    set_run(p.add_run(f"{title}. "), bold=True, color=NAVY)
    set_run(p.add_run(body))
    return p


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(value), size=9, color=WHITE, bold=True)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_width(cells[idx], widths[idx])
            set_cell_margins(cells[idx], top=95, bottom=95)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                set_cell_shading(cells[idx], PALE)
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            set_run(p.add_run(str(value)), size=font_size, color=INK)
    set_table_geometry(table, widths)
    keep_table_rows_together(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def page_break(doc):
    if not getattr(doc, "_cde_title_break_added", False):
        doc.add_page_break()
        doc._cde_title_break_added = True


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, color, before, after in (
        (1, 16, TEAL, 18, 10),
        (2, 13, TEAL, 14, 7),
        (3, 12, NAVY, 10, 5),
    ):
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.add_run("CYBERDECISIONENGINE  |  GUÍA DE LECTURA DE INFORMES"), size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("CyberDecisionEngine  ·  "), size=8.5, color=MUTED)
    add_field(footer, "PAGE")

    add_para(doc, "GUÍA PRÁCTICA", size=10, color=CYAN, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=72, after=16)
    add_para(doc, "Cómo leer y sustentar los informes", size=28, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "Informe ejecutivo e informe técnico", size=15, color=TEAL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=26)
    add_para(doc, "CyberDecisionEngine", size=18, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Ciberinteligencia estratégica para decisiones trazables", size=11, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=56)
    add_callout(doc, "Propósito", "Explicar, en lenguaje sencillo y defendible, de dónde sale cada dato, qué significa cada indicador, cómo se calculan los modelos y qué conclusiones sí o no pueden obtenerse.")
    add_para(doc, "Creado por Edwin Peñuela · Modelo desarrollado desde 2022", size=9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=22, after=3)
    add_para(doc, "Versión de la guía 1.0 · Julio de 2026", size=9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    page_break(doc)
    heading(doc, "1. La idea central en una página", 1)
    add_callout(doc, "Regla de oro", "El informe no empieza en una gráfica. Empieza en el alcance autorizado, continúa con registros trazables y solo termina en una decisión cuando la evidencia supera las reglas de validación.", "success")
    start_steps(doc)
    for title, body in (
        ("Alcance", "El usuario define marca, grupo o conglomerado, dominios propios, ventana temporal, sectores, regiones y, cuando aplique, comparativos separados."),
        ("Recolección", "Los conectores consultan fuentes públicas y autorizadas. Cada respuesta conserva origen, fecha, URL o referencia, estado y runId."),
        ("Procesamiento", "La plataforma normaliza fechas y URLs, resuelve entidades, deduplica y separa el registro original de su representación analítica."),
        ("Validación", "Un registro se vincula a una afirmación; se documentan método, relación, contradicciones, confianza y limitaciones."),
        ("Cálculo", "Los modelos usan únicamente variables disponibles y reglas versionadas. Sin cobertura suficiente se publica N/D, no un cero inventado."),
        ("Difusión", "Dashboard, informe ejecutivo, informe técnico, JSON y CSV usan el mismo RunContext. Cambia el nivel de detalle, no la fuente de verdad."),
    ):
        add_step(doc, title, body)
    heading(doc, "Tres distinciones que evitan errores", 2)
    add_bullet(doc, "Registro recolectado no equivale automáticamente a evidencia validada.")
    add_bullet(doc, "Intensidad de señales no equivale a probabilidad de ataque.")
    add_bullet(doc, "Sin datos no equivale a 0 %. Un cero requiere consulta exitosa, denominador válido y cobertura adecuada.")

    heading(doc, "2. De dónde vienen los datos", 1)
    add_table(
        doc,
        ["Capa", "Qué aporta", "Control de verdad"],
        [
            ("Entrada", "Organización, dominios propios, términos, sector, geografía y ventana.", "Queda registrada en el runId."),
            ("Fuentes públicas", "OSINT, noticias, advisories, DNS, RDAP/WHOIS, TLS, páginas e índices autorizados.", "URL/referencia, fecha, estado y fuente."),
            ("SOCMINT", "Menciones y relaciones públicamente indexadas cuando existen.", "No se dibujan nodos sin datos."),
            ("Dark Web", "Índices pasivos e importaciones autorizadas; no interacción con mercados.", "Metadatos minimizados y limitaciones visibles."),
            ("Catálogos", "CVE/CVSS/EPSS/KEV y mapeos ATT&CK, D3FEND, ATLAS, DISARM y controles.", "Versión y criterio de aplicabilidad."),
            ("Histórico", "Solo entra cuando se selecciona comparación o monitoreo continuo.", "Una corrida nueva no reutiliza silenciosamente evidencia anterior."),
        ],
        [1500, 4480, 3380],
        font_size=8.4,
    )
    add_callout(doc, "Cobertura operativa de conectores", "Se consulta en Configuración. Indica qué conectores fueron consultados, cuáles respondieron y cuántos registros entregaron. No mide riesgo, madurez ni ausencia de amenazas.")

    page_break(doc)
    heading(doc, "3. Vocabulario común", 1)
    add_table(
        doc,
        ["Término", "Cuándo se usa", "Qué no significa"],
        [
            ("Registro recolectado", "Dato recibido y normalizado.", "No confirma un hallazgo."),
            ("Evidencia validada", "Existe relación reproducible con una afirmación y método de validación.", "No necesariamente confirma un incidente."),
            ("Hallazgo validado", "Tiene estado validado/confirmado, método, fecha, validador y evidence_ids.", "No autoriza por sí solo una acción destructiva."),
            ("Alerta", "Regla, umbral, responsable, acción recomendada y estado.", "No es sinónimo de incidente."),
            ("Riesgo", "Combina plausibilidad contextual, impacto, controles y confianza.", "No es frecuencia observada."),
            ("SignalScore", "Concentración/intensidad de señales estratégicas trazables.", "No es probabilidad, cumplimiento ni madurez."),
            ("Presión validada", "Dirección estratégica publicada tras superar soporte y corroboración.", "No demuestra ataque o vulnerabilidad."),
            ("N/D", "Datos insuficientes o dimensión no calculable.", "No es 0 %."),
        ],
        [1900, 3940, 3520],
        font_size=8.8,
    )

    heading(doc, "4. Cómo leer el informe ejecutivo en 10 minutos", 1)
    add_callout(doc, "Audiencia", "Junta, CEO, comité de riesgos, CISO, líderes de fraude, legal, comunicaciones y negocio. Su función es decidir y priorizar; no sustituye el análisis técnico.")
    start_steps(doc)
    for title, body in (
        ("Portada y alcance", "Confirme organización, dominios, ventana, fecha del informe, runId y base de comparación. Si el alcance está mal, detenga la lectura."),
        ("Resumen ejecutivo", "Lea primero qué se encontró, nivel de confianza, limitaciones y la decisión que requiere atención."),
        ("Alertas validadas", "Verifique umbral, estado, responsable y acción. Una alerta sin esos elementos no debe tratarse como formal."),
        ("Riesgo y escenarios", "Observe riesgo inherente, residual, mapa 4x4 y bandas de sensibilidad. Pregunte qué evidencia soporta cada riesgo."),
        ("Cyber-PESTEL y Cyber-Porter", "Compare SignalScore, cobertura, confianza y presión validada. Revise los eventos que empujan o reducen cada dimensión."),
        ("Frameworks", "Úselos para orientar conversaciones de control. El porcentaje es cobertura de mapeo, no certificación ni cumplimiento."),
        ("Plan de trabajo", "Valide prioridad, horizonte, capacidad responsable, criterio de cierre y evidencia que justificó la acción."),
        ("Limitaciones", "Lea qué fuentes fallaron, qué no pudo validarse y qué no debe concluirse."),
    ):
        add_step(doc, title, body)
    add_callout(doc, "Frase útil para comité", "La conclusión está respaldada por registros trazables de esta corrida; la confianza y las limitaciones están separadas del impacto para evitar convertir cobertura en certeza.", "success")

    page_break(doc)
    heading(doc, "5. Cómo leer el informe técnico", 1)
    add_callout(doc, "Audiencia", "CISO, CTI, SOC, vulnerabilidades, infraestructura, fraude, GRC, forense, legal técnico y responsables de remediación.")
    start_steps(doc)
    for title, body in (
        ("Identidad de corrida", "Compruebe runId, versiones del motor/modelos, entradas, inicio, fin y ventana."),
        ("Superficie externa", "Revise dominios, subdominios, DNS, RDAP/WHOIS, TLS, tecnologías observables y puertos solo cuando exista recolección autorizada y verificable."),
        ("Pantallazos", "Una captura demuestra qué contenido era visible al momento de la recolección. No prueba por sí sola que la interpretación sea correcta; debe enlazarse a URL, fecha, hash y afirmación."),
        ("Vulnerabilidades", "Separe CVE aplicable de CVE meramente relacionada. CVSS es severidad técnica; EPSS estima explotación; KEV indica explotación conocida."),
        ("Evidencia URL por URL", "Revise consulta, respuesta disponible, canonical_url, hash, timestamps, fuente, entidad, relación, método, contradicciones y limitaciones."),
        ("PESTEL/Porter", "Audite clusters, fuentes independientes, relación con organización/sector/geografía y regla de publicación."),
        ("Mapeos y escenarios", "ATT&CK, D3FEND, ATLAS y DISARM clasifican o relacionan evidencia. Un mapeo no equivale a telemetría observada."),
        ("Cierre", "Cada hallazgo debe poder seguir la cadena afirmación → evidencia → interpretación → limitación → decisión → criterio de cierre."),
    ):
        add_step(doc, title, body)

    heading(doc, "6. Cómo interpretar porcentajes", 1)
    add_table(
        doc,
        ["Indicador", "Lectura correcta", "Error común"],
        [
            ("Cobertura de conectores", "Proporción/estado de conectores consultados y su respuesta operativa.", "Leerla como nivel de seguridad."),
            ("Cobertura de evidencia", "Masa de evidencia disponible para una dimensión.", "Asumir que 100 % significa verdad absoluta."),
            ("Confianza", "Calidad, diversidad, relación directa, acuerdo y extracción.", "Confundirla con impacto."),
            ("SignalScore", "Intensidad de señales relacionadas en escala 0–100.", "Llamarla probabilidad de ataque."),
            ("Presión validada", "Dirección de presión tras una puerta de publicación estricta.", "Tratarla como riesgo residual."),
            ("Riesgo residual", "Riesgo después de controles, con variables explícitas.", "Asumir que es una pérdida monetaria."),
            ("Mapeo de frameworks", "Cobertura de aspectos/control families relacionados.", "Presentarlo como cumplimiento o auditoría aprobada."),
        ],
        [2100, 4200, 3060],
        font_size=8.7,
    )

    page_break(doc)
    heading(doc, "7. Matemáticas explicadas sin perder rigor", 1)
    add_callout(doc, "Importante", "Las fórmulas son modelos de apoyo a la decisión. La plausibilidad contextual es heurística y versionada; no debe denominarse probabilidad calibrada si no existen métricas de calibración.", "warning")

    heading(doc, "7.1 Actividad de amenazas", 2)
    add_formula(doc, "T = Σᵢ (wᵢ · cᵢ · 2^(−edadᵢ / vidaMediaᵢ))", "Pondera fuente y confianza; reduce la influencia de señales antiguas.")
    add_formula(doc, "A = 1 − e^(−0.35 · T)", "Transforma volumen en una escala 0–1 y evita que una fuente muy prolífica domine sin límite.")

    heading(doc, "7.2 Plausibilidad contextual", 2)
    add_formula(doc, "L = 1 / (1 + e^(−z))", "La función logística acota el resultado entre 0 y 1.")
    add_formula(doc, "z = −2.10 + 0.70A + 0.85E + 0.75V + 0.90·logit(P)/6 + 0.85K + 0.70TTP + 0.55S + 0.35G − 0.80C − 0.60D − 0.45R", "A actividad; E exposición; V vulnerabilidad; P EPSS; K KEV; TTP técnicas; S sector; G geografía; C controles; D detección; R respuesta.")
    add_para(doc, "Lógica: las señales de amenaza/exposición elevan plausibilidad; controles, detección y respuesta la reducen. Los datos ausentes no deben convertirse automáticamente en cero observado.")

    heading(doc, "7.3 Impacto y controles", 2)
    add_formula(doc, "I = 0.25F + 0.20O + 0.20C + 0.15In + 0.10A + 0.05L + 0.05R", "Impacto financiero, operativo, confidencialidad, integridad, disponibilidad, legal y reputacional; los pesos suman 1.")
    add_formula(doc, "CE = 0.25ISO + 0.25NIST + 0.15SOC2 + 0.15D3FEND + 0.10Detección + 0.10Respuesta", "La efectividad de controles reduce el riesgo, con límite del 85 % para no prometer eliminación total.")

    heading(doc, "7.4 Riesgo inherente y residual", 2)
    add_formula(doc, "Rᵢₙₕ = 100 · L · I", "Combina plausibilidad e impacto en escala 0–100.")
    add_formula(doc, "Rᵣₑₛ = Rᵢₙₕ · [1 − min(0.85, CE)]", "Aplica el efecto de los controles sin reducir el riesgo a cero por construcción.")
    add_callout(doc, "Ejemplo didáctico", "Si L = 0.40, I = 0.60 y CE = 0.50, entonces R inherente = 24 y R residual = 12. Es un ejemplo matemático, no un hallazgo de una organización.")

    heading(doc, "7.5 Matriz 4×4 y sensibilidad", 2)
    add_formula(doc, "M = nivel(L) · nivel(I), con nivel ∈ {1,2,3,4}", "1–3 bajo; 4–7 medio; 8–11 alto; 12–16 crítico.")
    add_para(doc, "Monte Carlo toma muestras beta alrededor de L, I y CE con semilla reproducible y entrega p10, p50 y p90. Son bandas de sensibilidad del modelo, no una frecuencia observada ni una predicción temporal calibrada.")

    page_break(doc)
    heading(doc, "8. Cyber-PESTEL y Cyber-Porter", 1)
    add_callout(doc, "Qué analizan", "PESTEL organiza presiones cibernéticas macroambientales; Porter organiza fuerzas competitivas y exposición sectorial. Ambos se calculan para la marca, grupo o conglomerado y sus dominios propios, con contexto de sector, geografía, proveedores, clientes, sustitutos y competencia declarada.")
    add_para(doc, "Modelo productivo: strategic-evidence-v1.3.0. Las funciones heredadas pestel_cyber_index y porter_cyber_index no son el cálculo productivo actual.", size=9.5, color=MUTED, italic=True)

    heading(doc, "8.1 Calidad de cada contribución", 2)
    add_formula(doc, "qᵢ = 0.18Mᵢ + 0.16Qᵢ + 0.10Rᵢ + 0.16Dᵢ + 0.08Nᵢ + 0.12Cᵢ + 0.10Xᵢ + 0.10Gᵢ", "Combina ajuste temático, calidad, recencia, relación directa, novedad, corroboración, extracción y mapeo. Los coeficientes suman 1.")
    add_formula(doc, "m_d = Σᵢ (qᵢ · magnitudᵢ)", "Masa de evidencia de la dimensión d.")
    add_formula(doc, "cobertura_d = 1 − e^(−m_d / τ_d)", "La cobertura crece con evidencia útil y se satura de forma gradual.")

    heading(doc, "8.2 Intensidad, confianza y dirección", 2)
    add_formula(doc, "SignalScore_d = 100 · [0.65·cobertura_d + 0.20·directitud_d + 0.15·min(1, n_d/4)]", "Mide intensidad/concentración de señales, no riesgo ni probabilidad.")
    add_formula(doc, "Confianza_d = 100 · [0.20·cobertura_d + 0.80·(0.20·diversidad_d + 0.30·directitud_d + 0.20·acuerdo_d + 0.30·extracción_d)]", "Separa calidad analítica de intensidad.")
    add_formula(doc, "Presión_d = 50 + 50 · tanh(1.5 · z_d)", "z_d es la masa firmada dividida por la masa total. Solo se publica cuando el soporte supera la puerta de validación.")
    add_callout(doc, "Puerta por dimensión", "m_d ≥ 0.15 y, además, al menos dos clusters directos con dos fuentes independientes, o un evento oficial crítico validado.", "success")
    add_callout(doc, "Puerta agregada", "Cobertura calculable ≥ 60 %, confianza global ≥ 50 y más de la mitad de dimensiones calculables. Si no se cumple, se muestra N/D y se conservan aspectos/evidencias parciales.", "warning")

    heading(doc, "8.3 Dimensiones", 2)
    add_table(
        doc,
        ["Cyber-PESTEL", "Cyber-Porter"],
        [
            ("Geopolítica y política pública", "Rivalidad y presión competitiva"),
            ("Economía y sostenibilidad financiera", "Nuevos entrantes"),
            ("Factor humano y social", "Proveedores y terceros"),
            ("Tecnología y dependencia digital", "Clientes y canales"),
            ("Resiliencia y entorno operativo", "Sustitutos y cambios del mercado"),
            ("Legal y regulatorio", "Contexto sectorial transversal"),
        ],
        [4680, 4680],
    )

    page_break(doc)
    heading(doc, "9. Cómo sustentar el informe ante terceros", 1)
    add_para(doc, "Use esta secuencia de exposición. Mantiene la conversación en evidencia, decisión y límites.")
    start_steps(doc)
    for title, body in (
        ("Alcance", "Analizamos esta organización, estos dominios y esta ventana; los comparativos se mantienen separados."),
        ("Cobertura", "Estos conectores respondieron, estos fueron parciales y estas limitaciones condicionan la lectura."),
        ("Hechos", "Estos registros se recolectaron; estos pasaron validación; estos siguen como señales por revisar."),
        ("Interpretación", "El modelo explica intensidad, confianza, riesgo y contexto estratégico con métricas distintas."),
        ("Decisión", "Estas son posibilidades priorizadas, su responsable/capacidad y su criterio de cierre."),
        ("Límite", "No afirmamos ataque, incidente o cumplimiento cuando la evidencia no satisface el contrato semántico."),
    ):
        add_step(doc, title, body)
    heading(doc, "Preguntas que debe poder responder", 2)
    for item in (
        "¿Cuál es el runId y qué datos pertenecen exactamente a esa corrida?",
        "¿Qué demuestra la evidencia y qué no demuestra?",
        "¿Por qué una dimensión tiene porcentaje y otra aparece N/D?",
        "¿Qué variables elevaron el riesgo y cuáles lo redujeron?",
        "¿Qué fuente independiente corrobora el hallazgo?",
        "¿Cuál es el criterio de cierre y quién valida el resultado?",
    ):
        add_bullet(doc, item)

    heading(doc, "10. Errores de lectura frecuentes", 1)
    add_table(
        doc,
        ["Error", "Corrección"],
        [
            ("“El radar dice 70 %, entonces hay 70 % de probabilidad de ataque.”", "Verifique qué indicador es: intensidad, riesgo o confianza. Solo llame probabilidad a un modelo calibrado."),
            ("“No hubo datos, entonces el valor es cero.”", "Use N/D. Cero exige consulta exitosa, cobertura adecuada y denominador válido."),
            ("“La URL existe, por tanto el hallazgo es crítico.”", "La criticidad requiere relación, condición de seguridad, impacto y validación; una URL por sí sola no basta."),
            ("“El mapeo NIST/ISO muestra cumplimiento.”", "Muestra cobertura de mapeo de controles, no certificación ni auditoría de cumplimiento."),
            ("“ATT&CK aparece, entonces la técnica fue observada.”", "ATT&CK observado exige telemetría adversaria, comportamiento, timestamp, activo y evidence_ids."),
            ("“La captura confirma la interpretación.”", "La captura confirma contenido visible; la interpretación debe validarse aparte."),
        ],
        [3920, 5440],
        font_size=8.2,
    )

    page_break(doc)
    heading(doc, "11. Checklist de revisión", 1)
    heading(doc, "Antes de aprobar el informe ejecutivo", 2)
    for item in (
        "Alcance, dominios, fecha, ventana y runId son correctos.",
        "Cada alerta tiene regla, umbral, responsable, acción y estado.",
        "Los porcentajes indican claramente qué miden y qué no miden.",
        "PESTEL/Porter muestran cobertura, confianza, intensidad y presión por separado.",
        "Las recomendaciones corresponden a hallazgos o se marcan como preventivas.",
        "Las limitaciones y conectores parciales son visibles.",
    ):
        add_bullet(doc, "☐ " + item)
    heading(doc, "Antes de aprobar el informe técnico", 2)
    for item in (
        "Toda afirmación importante enlaza evidence_ids.",
        "Toda evidencia tiene fuente, URL/referencia, fecha, hash y método.",
        "Los dominios/subdominios pertenecen al alcance o se etiquetan como externos.",
        "CVE/EPSS/KEV distinguen aplicabilidad de referencia general.",
        "Pantallazos y URLs no están solapados y se pueden ampliar/abrir.",
        "No hay secretos, rutas locales ni datos de otra organización.",
        "JSON, CSV y HTML reflejan la misma fuente de verdad.",
    ):
        add_bullet(doc, "☐ " + item)

    doc.add_page_break()
    heading(doc, "12. Glosario mínimo", 1)
    add_table(
        doc,
        ["Concepto", "Definición breve"],
        [
            ("runId", "Identificador único que enlaza parámetros, recolección, análisis, dashboard, exportes e informes."),
            ("RunContext", "Fuente de verdad estructurada de una corrida."),
            ("Claim", "Afirmación que debe ser respaldada, limitada y decidida."),
            ("Evidence", "Registro con procedencia y relación explícita con una afirmación."),
            ("Confidence", "Calidad del soporte; no es impacto."),
            ("TLP/PAP", "Reglas de distribución y acción permitida sobre la información."),
            ("Deduplicación", "Elimina repeticiones exactas o equivalentes sin borrar el registro original."),
            ("Corroboración", "Apoyo de fuentes independientes para la misma afirmación."),
            ("N/D", "No disponible/no calculable con la cobertura actual."),
        ],
        [1900, 7460],
    )
    add_callout(doc, "Cierre", "La calidad del informe no se mide por cuántos gráficos contiene, sino por la capacidad de explicar cada afirmación, mostrar su evidencia, reconocer sus límites y convertirla en una decisión verificable.", "success")
    add_para(doc, "CyberDecisionEngine · Creado por Edwin Peñuela · 2022–2026", size=9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=0)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
